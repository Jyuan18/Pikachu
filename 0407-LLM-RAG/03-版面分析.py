# 富文本txt读取
# read()
with open("./text.txt", 'r', encoding='utf-8') as f:
    data = f.read()
    print(data)

# readline()
with open('./text.txt', 'r', encoding='utf-8') as f:
    data = f.readline()
    print(data)

# readlines()
with open('./text.txt', 'r', encoding='utf-8') as f:
    data = f.readlines()
    print(data)

with open('./text.txt', 'r', encoding='utf-8') as f:
    for ann in f.readlines():
        ann = ann.strip('\n')
        print(ann)

# ==========================================================================================================

# pdf解析
# pdfplumber
# 提取文本
import pdfplumber
file_name = '**.pdf'
output_file = '**.txt'
with pdfplumber.open(file_name) as p:
    page_count = len(p.pages)
    for i in range(0, page_count):
        page = p.pages[i]
        textdata = page.extract_text()
        data = open(output_file, 'a')
        data.write(textdata)

# 提取表格
import pdfplumber
file_name = '**.pdf'
with pdfplumber.open(file_name) as p:
    page_count = len(p.pages)
    for i in range(0, page_count):
        page = p.pages[i]
        textdata = page.extract_table()
        data = open(output_file, 'a')
        data.write(textdata)

# 提取表格文本,保存为excel文件
import pdfplumber
from openpyxl import Workbook
file_name = '**.pdf'
output_file = '**.xlsx'
with pdfplumber.open(file_name) as pdf:
    page01 = pdf.pages[0]
    table = page01.extract_table()
    workbook = Workbook()
    sheet = workbook.active
    for row in table:
        sheet.append(row)
    workbook.save(filename=output_file)

# 提取图片,保存本地
import pdfplumber
file_name = '**.pdf'
output_file = '**.xlsx'
with pdfplumber.open(file_name) as pdf:
    # 获取第一页
    first_page = pdf.pages[1]
    print('页码:', first_page.page_number)
    print('页宽:', first_page.width)
    print('页高:', first_page.height)
    text = first_page.extract_text()
    print(text)
    # 获取第一页的图片,获取的是一个列表,列表中存储的是字典
    imgs = first_page.images
    i = 0
    for img in imgs:
        print(img['stream'].get_data())
        with open(output_file, model='wb') as f2:
            f2.write(img['stream'].get_data())

# ============================================================
# PyMuPDF
# Fitz是PyMuPDF的子模块,提供了一个简化和封装版本的PyMuPDF功能
import fitz
pdf_path = ''
doc = fitz.open(pdf_path)
# base info
title = doc.metadata['title']
author = doc.metadata['author']
create_data = doc.metadata['creationDate']
num_pages = doc.page_count
page = doc.load_page(0)
page_height = page.bound().height
page_width = page.bound().width
# 获取文本
for page_idnex in range(num_pages):
    page = doc.load_page(page_idnex)
    text = page.get_text()
# 获取图片
for page_index in range(num_pages):
    page = doc.load_page(page_index)
    image_list = page.get_images()
    for img in image_list:
        xref = img[0]
        pix = fitz.Pixmap(doc, xref)
        print(pix.colorspace, '-->', fitz.csRGB)
        img_path = f'./output/image{page_index+1}_{xref}.png'
        pix.save(img_path)
# 获取表格
for page_index in range(num_pages):
    page = doc.load_page(page_index)
    tables = page.find_tables()
    print(f'tables: {tables}')
    for i, table in enumerate(tables):
        df = tables[0].to_pandas()
        print(df.head())
        df.to_csv(f'../output/table_pg_{page_index}_{i+1}.csv', index=False)
# pdf文档分割
import fitz
pdf_path = '../test.pdf'
pdf_document = fitz.open(pdf_path)
num_pages = pdf_document.page_count
# 构建输出文件名,以页数命名
output_pdf = '../output/book_split_{0}.pdf'
for i in range(i, num_pages):
    print(f'i:{i}')
    new_pdf = fitz.open()
    new_pdf.insert_pdf(pdf_document, from_page=i-1, to_page=i)
    new_pdf.save(output_pdf.format(i))
    new_pdf.close()
    pdf_document.close()

# =========================================================================
# 如何借助大模型进行文档回答
import os
import fitz
from openai import OpenAI

def get_pdf_content(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    num_pages = doc.page_count
    bg_content_list = []
    for page_index in range(num_pages):
        page = doc.load_page(page_index)
        text = page.get_text()
        bg_content_list.append(text)
    return ''.join(bg_content_list)

def get_answer(pdf_content: str, query: str) -> str:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    response = client.chat.completions.create(
        model='gpt-3.5-turbo',
        messages=[
            {'role': 'system', 'content': 'You are a helpful assistant.'},
            {'role': 'user', 'content': f'The full text of PDF file is:{pdf_content}'},
            {'role': 'user', 'content': query}
        ],
        max_tokens=1000
    )
    answer = response.choices[0].message.content
    return answer

if __name__ == '__main__':
    content = get_pdf_content('../test.pdf')
    query1 = 'OPPO Find N3 Flip 的价格?'
    print(get_answer(pdf_content=content, query=query1))

    query2 = '蚂蚁集团发布的大模型叫什么?'
    print(get_answer(pdf_content=content, query=query2))

    query3 = '混元大模型是什么时候发布的?'
    print(get_answer(pdf_content=content, query=query3))

# ==========================================================================================================
# BeautifulSoup4-bs4
from bs4 import BeautifulSoup
html = """
<html><head><title>The Dormouse's story</title></head>
<body>
<p class="title"><b>The Dormouse's story</b></p>
<p class="story">once upon a time there were three little sisters; and their names were
<a href='http://example.com/elsie' class='sister' id='link1'><!--Elsie--></a>,


"""
# beautifulsoup对象
soup = BeautifulSoup(html, 'lxml')
print(f"type(soup):{type(soup)} \n")
# tag对象
print(f'soup.head:{soup.head} \n')
print(f'soup.head.name:{soup.head.name} \n')
print(f'soup.head.attrs:{soup.head.attrs} \n')
print(f'type(soup.head):{type(soup.head)} \n')
print()
# navigable string对象
print(f'soup.title.string:{soup.title.string} \n')
print(f'type(soup.title.string):{type(soup.title.string)} \n')
# comment对象
print(f'soup.a.string:{soup.a.string} \n')
print(f'type(soup.a.string):{type(soup.a.string)} \n')
# 结构化输出soup对象
print(f'soup.prettify()=>{soup.prettify()}')

# TODO
# 未学习完
# https://articles.zsxq.com/id_ptqivfoinkp8.html

# ==========================================================================================================

# docx
# pip install python-docx
from docx import Document
file_name = './test.docx'
doc = Document(file_name)
print(f'\n遍历文档的段落:')
for para in doc.paragraphs:
    print(para.text)
for table in doc.tables:
    row_count = len(table.rows)
    col_count = len(table.columns)
    for i in range(row_count):
        row  = table.rows[i].cells
        print(f'row:{row}')
    for i in range(row_count):
        for j in range(col_count):
            print(table.cell(i,j).text)